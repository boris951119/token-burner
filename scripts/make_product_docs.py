# -*- coding: utf-8 -*-
r"""生成产品文档：产品介绍 + 竞品对比分析报告（均为 Word）。

运行：python scripts\make_product_docs.py
输出：token-burner_产品介绍.docx
      token-burner_竞品对比分析报告.docx
"""
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent / "_docgen"))

from docx import Document  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402
from docx.shared import Pt  # noqa: E402

ROOT = Path(__file__).resolve().parent.parent
FONT = "微软雅黑"


def new_doc():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    return doc


def h(doc, level, text):
    p = doc.add_heading(text, level)
    for r in p.runs:
        r.font.name = FONT
        r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    return p


def para(doc, text, bold=False, size=11, color=None, align=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = FONT
    r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    r.font.size = Pt(size)
    r.font.bold = bold
    if align:
        p.alignment = align
    return p


def bullet(doc, text, bold_head=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_head:
        r = p.add_run(bold_head)
        r.font.name = FONT
        r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        r.font.size = Pt(11)
        r.font.bold = True
    r = p.add_run(text)
    r.font.name = FONT
    r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    r.font.size = Pt(11)
    return p


def table(doc, rows, widths=None, size=9.5):
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = "Table Grid"
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.cell(ri, ci)
            cell.text = ""
            p = cell.paragraphs[0]
            r = p.add_run(val)
            r.font.name = FONT
            r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
            r.font.size = Pt(size)
            r.font.bold = ri == 0
    if widths:
        for ci, w in enumerate(widths):
            for ri in range(len(rows)):
                t.cell(ri, ci).width = w
    doc.add_paragraph()
    return t


def cover(doc, subtitle, tag):
    for _ in range(5):
        doc.add_paragraph()
    p = para(doc, "token-burner · 文档消耗器", bold=True, size=28,
             align=WD_ALIGN_PARAGRAPH.CENTER)
    p = para(doc, "AI 多智能体项目团队系统", size=20, align=WD_ALIGN_PARAGRAPH.CENTER)
    p = para(doc, subtitle, bold=True, size=24, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    para(doc, tag, size=13, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()


# =========================================================================
# 文档一：产品介绍
# =========================================================================
doc = new_doc()
cover(doc, "产品介绍", "（v0.5.0-beta · 2026 年 9 月）")

h(doc, 1, "一、产品概述")
para(doc, "token-burner（文档消耗器）是一套 AI 多智能体项目团队系统：用户用一句自然"
     "语言描述软件需求，系统自动完成「评估路由 → 多模型组队 → 方案讨论 → spec 确认 → "
     "模块拆分 → 逐模块写码/测试/执行 → 交付物汇总」的完整开发闭环，最终产出一个结构"
     "完整、带测试与文档、可本地运行的软件项目。")
para(doc, "产品的核心理念可以概括为一句话：决策归 LLM，校验与边界归程序。大模型负责"
     "需求理解、方案设计与代码生成这类需要智能的判断；而成本控制、接口纪律、执行安全、"
     "流程收敛这些不允许「靠自觉」的环节，全部由确定性的程序机制把关。")
para(doc, "另一条贯穿全局的设计哲学是「Set your tokens on fire——让每个 token 都烧在"
     "刀刃上」：多智能体协作天然放大 token 消耗，token-burner 用六层确定性护栏把这种"
     "消耗从「不可预估的账单风险」变成「开工前即可预算、事后可逐条审计的工程量」。")
para(doc, "当前版本 v0.5.0-beta：833 项自动化测试全量通过，提供 Web 工作台、桌面客户端、"
     "命令行、VS Code 插件四种入口，并具备 GitHub Actions 自动发布流水线。", bold=True)

h(doc, 1, "二、产品要解决的问题")
bullet(doc, "LLM 多轮对话与多模型协作让账单在无声中膨胀；讨论发散、修复死循环、"
       "单次输出倾泻等场景下，企业无法对 AI 开发做预算管理。",
       "成本失控：")
bullet(doc, "单一模型「自己写、自己审」缺乏独立视角，盲区一致、幻觉互相印证；"
       "模型市场碎片化的今天，任何单一模型都难以在架构、实现、测试三个视角同时占优。",
       "能力天花板：")
bullet(doc, "主流工具生成的代码缺乏接口契约、静态验证与执行反馈闭环，「看起来对」"
       "不等于「跑得对」，难以直接进入生产链路。", "质量无闸门：")
bullet(doc, "AI 开发过程是黑盒：用了多少 token、哪一步花掉了预算、为什么修复了五轮——"
       "事后无法回答，企业合规与审计无从谈起。", "过程不可审计：")
para(doc, "破局思路：多智能体异构协作解决质量问题，确定性护栏体系解决成本问题，"
     "全链路落盘解决审计问题——三者缺一不可，这正是 token-burner 的产品定义。", bold=True)

h(doc, 1, "三、产品形态：一个内核，四个入口")
table(doc, [
    ("入口", "使用方式", "适合人群"),
    ("Web 工作台", "python -m app.server 后浏览器访问 127.0.0.1:8000；"
     "实时监控、对话流图、成本看板、设置页", "最完整的可视化体验，推荐首选"),
    ("桌面客户端", "独立 exe 双击即用（约 75MB，无需安装 Python）；"
     "单进程零网络架构，兼容企业安全环境", "非技术用户、最小白的使用方式"),
    ("命令行 CLI", "python -m app.main 交互式引导（评估展示 → 选模型 → 选模式 → 交付）",
     "脚本化、自动化、CI 集成场景"),
    ("VS Code 插件", "活动栏面板发起任务、SSE 实时进度、生成代码 diff 预览与一键应用、"
     "成本统计、历史任务与一键续跑", "在 IDE 内完成全部流程的开发者"),
], widths=[1.1, 3.3, 2.6])
para(doc, "四种入口共用同一套编排内核与 API 契约：桌面端面向小白，CLI 面向自动化，"
     "Web 工作台与 VS Code 插件面向日常开发者。")

h(doc, 1, "四、核心功能详解")

h(doc, 2, "4.1 需求评估与智能路由")
bullet(doc, "评估主 LLM 对需求打难度分（1-10）并三分类：基础任务直答（零团队成本）、"
       "研究·分析任务深度直出、编程任务进入完整团队流程；难度 ≤3 的简单编程直接单文件"
       "直出，不组建团队（省 token 节流）。", "三分类路由：")
bullet(doc, "System-1 快判 + System-2 全量评估双通道（可选开启）：轻量模型先以极低成本"
       "识别闲聊/无意义输入直接拒答，编程或低置信度请求自动升级全量评估——"
       "「宁升勿误」，省 token 不牺牲准确性。", "双模式意图识别：")
bullet(doc, "按难度分三档分发（可选开启）：难度 ≥7 全旗舰模型、4-6 主力+轻量混合、"
       "<4 全轻量；档位列表与阈值均可在 config.json 自定义，成本看板同步输出"
       "「实际成本 vs 全旗舰假设成本」对比。", "智能模型路由：")
bullet(doc, "边界复核护栏对评估结果做二次校验，不确定时保守视作编程任务并向用户确认，"
       "避免误判导致流程走错。", "保守降级确认：")

h(doc, 2, "4.2 多智能体团队协作")
bullet(doc, "主 LLM（架构师/PM）+ 开发副 LLM + 测试副 LLM，三者模型强制互异——"
       "独立视角交叉验证，从结构上杜绝同源盲区；支持接入 OpenAI、Anthropic、DeepSeek、"
       "Gemini、智谱 GLM 等任意 litellm 兼容模型。", "固定团队编制：")
bullet(doc, "主 LLM 提出初始方案 → 开发/测试副 LLM 双评审打分（评分 ≥8 的建议自动采纳）"
       "→ 主 LLM 汇总修订 → 收敛生成 spec.md；默认最多 3 轮，讨论循环（Jaccard 字面"
       "重复 + embedding 语义重复双道检测）达到阈值即冻结副 LLM 发言权，由主 LLM "
       "收权裁决。", "方案讨论流程：")
bullet(doc, "用户对 spec 的修改意见默认最多 3 轮，之后主 LLM 主动收敛合并输出最终"
       " spec，不再反复征询。", "spec 确认收敛：")
bullet(doc, "开发前对陌生技术栈生成四段式结构化摘要（来源/版本/用法示例/坑点）注入"
       "开发与测试提示词；支持联网搜索（duckduckgo 免密钥 / tavily），拥有独立预算"
       "（默认 20k token）与 SQLite 缓存（TTL 7 天）。", "Researcher 前置调研（可选）：")

h(doc, 2, "4.3 模块化工程与接口门禁")
bullet(doc, "难度 ≥5 或预估文件数 ≥6 的任务自动拆分为多个功能模块，每个模块拥有"
       "独立目录、独立规格 .md 与独立修复记录——支持单独修改、单独验收、单独追踪。",
       "条件化模块拆分：")
bullet(doc, "拆分时生成接口地图（imports/exports/public_api）作为模块间唯一合法"
       "基线；AST 接口门禁对生成代码做引用抽取与差异判定，违规引用直接阻断进入修复"
       "循环——零执行、零 LLM 调用，是全系统唯一不消耗 token 的纪律执行器。",
       "接口契约与 AST 门禁：")
bullet(doc, "公共共享库（_shared）发生变更时触发整包回归，保障「改 A 不动 B」的模块"
       "隔离承诺。", "共享库整包回归：")
bullet(doc, "生成项目自动 git 初始化并按阶段本地提交，交付即带完整版本历史。",
       "Git 版本管理：")

h(doc, 2, "4.4 双模执行验证")
table(doc, [
    ("维度", "安全审阅模式（默认）", "自动验证模式"),
    ("执行方式", "不执行任何生成代码；AST 静态检查 + LLM 逻辑审查",
     "危险预扫描后真实执行生成代码与测试"),
    ("结果反馈", "用户本地运行后反馈：成功确认 / 报错修复 / exit 停止",
     "结构化执行结果自动注入修复循环"),
    ("修复上限", "5 轮，超限输出「已知问题与降级方案」", "同左"),
    ("任务预算", "200k token", "×2.5 = 500k token"),
    ("隔离级别", "无执行，天然安全", "Docker 容器沙箱；Docker 缺失自动降级进程模式"),
], widths=[1.0, 2.9, 2.9])
bullet(doc, "内存 512MB / CPU 1.0 / 进程数 128 / tmpfs 64MB / 只读文件系统 / 无网络 / "
       "非 root 运行 / 30 秒超时熔断；超限由内核直接终止（exit 137 判定失败）。",
       "容器沙箱资源配额：")
bullet(doc, "Python（python:3.11-slim + pytest）与 Node.js（node:20-slim + "
       "node --test，TypeScript 支持预热）；镜像自动预热与可用性检测。", "多语言运行时：")

h(doc, 2, "4.5 六层成本护栏体系")
table(doc, [
    ("层", "护栏", "默认参数", "达标行为"),
    ("0", "单任务总预算（总闸）", "200k tokens · 自动模式 ×2.5",
     "≥90% 进入省 token 模式；超预算中止并落盘已完成部分与止损清单"),
    ("1", "讨论轮数上限", "3 轮", "主 LLM 直接产出收敛 spec"),
    ("2", "单轮输出上限", "8k tokens", "截断按未完成处理，分块续写（≤2 次）"),
    ("3", "讨论循环检测", "Jaccard 0.9 / embedding 0.85 / 重复 3 次",
     "冻结副 LLM 发言权，主 LLM 收权裁决"),
    ("4", "修复循环上限", "5 轮", "输出「已知问题与降级方案」交用户决策"),
    ("5", "spec 确认收敛", "3 次", "合并意见输出最终 spec"),
], widths=[0.4, 1.7, 2.0, 2.9])
para(doc, "设计原则：所有护栏由程序确定性执行（AST 解析、相似度计算、预算累加），"
     "不依赖 LLM 自律。任务开工前，成本上界即可预估；任务结束后，logs/cost_report.json "
     "逐条记录每次调用的 token 消耗——预算工程而非事后账单。", bold=True)

h(doc, 2, "4.6 可视化与运维")
bullet(doc, "阶段耗时条形图、token 累积曲线（实时 SVG）、模块状态全景；"
       "异步任务 API（提交即返回 task_id）+ SSE 事件流实时推送。", "实时监控面板：")
bullet(doc, "讨论消息结构化记录并以纵向节点链渲染，AI 团队「怎么吵出来的」全程可见。",
       "Agent 对话流图：")
bullet(doc, "基于历史项目统计（需求词袋相似度 → 成功率优先 → 平均成本 → 保守缺省）"
       "推荐执行模式与预算档位——纯确定性统计，零 LLM 参与，用户可覆盖。",
       "模式智能推荐：")
bullet(doc, "按模型/阶段/档位多维统计、实际成本 vs 旗舰假设成本对比、Embedding 与 "
       "Researcher 缓存节省量与命中率。", "成本看板：")
bullet(doc, "深/浅双主题（跟随系统实时切换）、分类设置页（外观/模型/预算/安全四组）、"
       "Web 与桌面端同源。", "界面体验：")

h(doc, 2, "4.7 可靠性工程")
bullet(doc, "每个任务的状态快照落盘（sessions/pipeline_state.json），进程崩溃或"
       "中断后重启即恢复，已完成模块自动跳过。", "中断恢复：")
bullet(doc, "运行中任务协作式取消（取消标志注入预算检查点）、服务重启后自动清扫"
       "僵尸任务。", "任务取消与僵尸清理：")
bullet(doc, "单次 LLM 调用 120 秒超时 + 瞬态错误指数退避重试（≤3 次）；按供应商"
       "令牌桶全局限流（超限排队不报错）。", "调用韧性：")
bullet(doc, "ModelClient 任务级隔离（每任务独立实例）+ 项目级锁管理，支持 ≥4 路任务"
       "并发且互不污染。", "并发架构：")
bullet(doc, "Embedding 缓存（内容哈希键，7 天过期）、讨论论点库持久化、"
       "Researcher 调研缓存——增量比对压掉重复计算成本。", "三级缓存：")

h(doc, 1, "五、一次任务的完整旅程（端到端流程）")
steps = [
    ("1. 点火", "用户在任一入口输入需求文本（自动包裹数据边界，防提示词注入）。"),
    ("2. 评估", "（可选快判拦截闲聊后）评估主 LLM 输出难度分与三分类结论；"
     "简单编程直接单文件直出，流程结束。"),
    ("3. 组队", "用户确认三模型（强制互异）与执行模式（auto 模式二次确认）；"
     "（可选）Researcher 先行调研注入领域知识。"),
    ("4. 讨论", "主 LLM 初始方案 → 双评审 → 修订循环（护栏约束下最多 3 轮），"
     "全程对话流可回放。"),
    ("5. spec", "生成 spec.md 请用户确认（≤3 轮意见合并后强制收敛）。"),
    ("6. 拆分", "复杂任务按条件模块化，生成接口地图与接口契约。"),
    ("7. 开发循环", "逐模块「写码 → 写测试 → 静态验证 → AST 接口门禁 → 执行验证 → "
     "修复（≤5 轮）」，安全模式每模块征询用户反馈。"),
    ("8. 交付", "汇总交付物与成本报告；超预算则落盘止损清单交用户决策。"),
    ("9. 恢复", "任何时刻中断均可续跑——已完成模块自动跳过，反馈可继续驱动修复闭环。"),
]
table(doc, [("阶段", "发生什么")] + steps, widths=[1.1, 5.9])

h(doc, 1, "六、交付物说明")
para(doc, "每个任务产出独立项目目录 projects/<需求>_<时间戳>/：")
table(doc, [
    ("内容", "说明"),
    ("modules/*.md", "每个模块的独立规格文档（可按「功能→规格→代码→变更记录」回溯）"),
    ("code/", "模块化源代码（含项目级 conftest 与包结构，开箱可运行）"),
    ("tests/", "每个模块的独立测试文件（附带运行指令与预期输出）"),
    ("changelog/", "历次修复记录"),
    ("sessions/", "需求原文、spec、接口地图、方案讨论记录、任务/管线状态快照、"
     "Researcher 调研摘要"),
    ("logs/cost_report.json", "成本审计报告：逐调用 token、按模型/阶段/档位统计、"
     "节省量"),
    ("（git 仓库）", "阶段性本地提交的完整版本历史"),
], widths=[2.0, 5.0])

h(doc, 1, "七、安全与信任边界")
bullet(doc, "用户需求、运行反馈、报错日志、LLM 生成代码、Researcher 摘要与联网抓取"
       "文本，全部按不可信输入处理：数据边界包裹 + 提示词明示「是数据非指令」+ 截断"
       "限长。", "不可信输入治理：")
bullet(doc, "自动模式默认 Docker 容器级隔离（资源配额/只读/无网络/非 root/超时熔断），"
       "降级进程模式仍有危险 API 黑名单预扫描兜底；安全模式不执行任何生成代码。",
       "执行隔离：")
bullet(doc, "API 密钥只从环境变量读取，不写入代码与配置文件。", "密钥管理：")

h(doc, 1, "八、技术架构")
para(doc, "系统自上而下分四层：")
table(doc, [
    ("层级", "组成", "职责"),
    ("交互层", "Web 工作台 / 桌面客户端 / CLI / VS Code 插件",
     "统一 API 契约，交互层零流程逻辑"),
    ("编排层", "orchestrator（路由/组队/讨论）、pipeline（管线/恢复）、"
     "task_manager（异步任务/SSE/取消）、recommender（模式推荐）",
     "流程推进与状态管理"),
    ("智能体层", "dev_loop（开发循环）、module_builder（模块构建）、"
     "researcher + web_research（前置调研）", "各角色的执行逻辑"),
    ("基础设施层", "model_client（litellm 统一接入/重试/限流）、budget（预算闸门）、"
     "similarity（循环检测）、interface_check（AST 门禁）、docker/local executor"
     "（双执行器）、cost_dashboard（成本看板）",
     "确定性校验与资源治理"),
], widths=[0.9, 3.1, 3.0])

h(doc, 1, "九、质量与工程成熟度")
bullet(doc, "833 项自动化测试全量通过（全 stub 设计，无需 API 密钥即可回归），"
       "含 4 路并发压测与消息协议契约测试。", "测试体系：")
bullet(doc, "快慢双模式 A/B 评测框架（token/误判率/延迟三维对比）、性能基线脚本"
       "（标准任务集可复现对比）。", "评测框架：")
bullet(doc, "GitHub Actions：打标签自动触发全量回归 → PyInstaller 构建 → 产物体积"
       "检查（≤80MB）→ Release 附 EXE。", "发布流水线：")
bullet(doc, "规格驱动（SDD）+ 测试驱动（TDD）开发方法，全部能力带规格锚点与回归"
       "安全网。", "开发方法论：")

h(doc, 1, "十、配置与扩展")
bullet(doc, "模型无关：基于 litellm，OpenAI / Anthropic / DeepSeek / Gemini / 智谱 GLM "
       "等任意兼容端点自由组合，config.json 一行切换。", "模型接入：")
bullet(doc, "六层护栏阈值、路由档位、Docker 配额、Researcher 预算等 40+ 参数集中"
       "config.json 管理，非法配置尽早失败。", "参数化：")
bullet(doc, "所有 v0.4/v0.5 新能力缺省关闭，默认行为与 v0.3.1 完全一致——"
       "按需灰度启用，升级零迁移。", "渐进式启用：")

out1 = ROOT / "token-burner_产品介绍.docx"
doc.save(out1)
print("saved:", out1)


# =========================================================================
# 文档二：竞品对比分析报告
# =========================================================================
doc = new_doc()
cover(doc, "类似产品对比分析报告", "（v0.5.0-beta · 2026 年 9 月 · 信息时点见文末来源）")

h(doc, 1, "一、报告概述")
para(doc, "本报告将 token-burner 置于 2026 年 AI 编码工具市场中，与三类主要玩家——"
     "自主 AI 软件工程师（Devin 等）、多智能体开发框架（MetaGPT / ChatDev 等）、"
     "AI 编码助手（GitHub Copilot / Cursor / Claude Code 等）——进行系统对比，"
     "客观分析差异化定位、竞争优势与真实差距，为产品决策与对外沟通提供依据。")
para(doc, "对比维度覆盖：产品形态、定价模型、多智能体架构、成本治理、验证闭环、"
     "沙箱隔离、模型中立性、可审计性与工程成熟度。")

h(doc, 1, "二、市场格局：三类玩家")
table(doc, [
    ("类别", "代表产品", "核心特征", "与 token-burner 的关系"),
    ("A. 自主 AI 软件工程师", "Devin（Cognition）、OpenHands",
     "云端托管的全自主工程师，沙箱 VM 内规划-编码-测试-交付 PR",
     "定位最接近的直接竞品；但为闭源 SaaS，按 ACU 后付费"),
    ("B. 多智能体开发框架", "MetaGPT、ChatDev、AutoGen、CrewAI",
     "开源框架，角色化多智能体协作生成软件，偏研究与自建",
     "技术路线同源；token-burner 是该路线的产品化与工程化"),
    ("C. AI 编码助手", "GitHub Copilot、Cursor、Claude Code、Windsurf、Aider",
     "嵌入 IDE/终端，辅助人类开发者（补全/对话/局部代理）",
     "错位竞争：助手服务「写代码的人」，token-burner 服务「要项目的人」"),
], widths=[1.4, 1.7, 2.0, 1.9])
para(doc, "token-burner 横跨 A、B 两类：采用 B 类的多智能体编排路线，但以 A 类的"
     "「输入需求、输出可运行项目」为交付目标，并以本地私有化 + 模型中立（自带密钥）"
     "区别于云端 SaaS。", bold=True)

h(doc, 1, "三、重点竞品逐项分析")

h(doc, 2, "3.1 Devin（Cognition）——定位最接近的直接竞品")
para(doc, "概况：全球首个全自主 AI 软件工程师。用户以自然语言派发任务，Devin 在沙箱"
     "虚拟机（浏览器+终端+编辑器）中规划、编码、测试并产出 GitHub/GitLab PR；2026 年"
     "已支持并行实例、AI 代码评审（Devin Review）、交互式规划，并通过 FedRAMP 高等级"
     "授权进入美国政府市场，企业客户含 Goldman Sachs、Dell、Mercedes-Benz 等。")
para(doc, "定价：Core 版 $20/月起 + 按 ACU（Agent Compute Unit）计量，$2.25/ACU，"
     "1 ACU ≈ 15 分钟计算；中等复杂任务消耗 5-20 ACU，即单任务真实成本 $11-45，"
     "团队月度实际账单常达 $200-2,250。")
para(doc, "关键实测数据（独立测试）：有明确复现路径的 bug 修复成功率约 78%；"
     "但跨多文件、含复杂依赖的真实工程任务成功率仅 14-15%；涉及 50+ 文件的任务"
     "仍有可观失败率。")
table(doc, [
    ("Devin 优势", "Devin 局限（token-burner 的机会）"),
    ("全自主端到端交付，PR 级集成", "闭源 SaaS：代码与过程数据出域，数据敏感企业受限"),
    ("企业级治理（VPC 部署/审计/权限）", "ACU 后付费：成本开工前不可预估，无 token 级预算闸门"),
    ("品牌与生态：Fortune 50 渗透、FedRAMP", "模型黑盒路由：厂商锁定，无法自选三模型互验"),
    ("并行实例、Dev Review 等成熟周边", "复杂任务 14-15% 成功率说明编排层质量闸门仍有空间"),
], widths=[3.5, 3.5])

h(doc, 2, "3.2 MetaGPT（DeepWisdom）——技术路线同源的开源标杆")
para(doc, "概况：开源多智能体框架（MIT 协议，GitHub 69,000+ star，ICLR 2024 Oral），"
     "以「Code = SOP(Team)」为理念，用 PM/架构师/项目经理/工程师/QA 五个角色化智能体"
     "沿标准作业流程协作，从一行需求产出 PRD、设计文档、任务列表、源码与测试；"
     "已商业化為 MGX（2026 年更名 Atoms）。")
table(doc, [
    ("MetaGPT 优势", "MetaGPT 局限（token-burner 的机会）"),
    ("学术影响力大、SOP 结构化协作理念成熟", "研究框架属性：无预算闸门与成本看板，token 消耗不可治理"),
    ("完全开源，社区活跃", "无产品化交付形态：无桌面端/插件/异步任务/中断恢复等工程能力"),
    ("中间产物（PRD/设计）结构完整", "无确定性接口门禁与沙箱执行，质量约束靠提示词而非机制"),
], widths=[3.5, 3.5])

h(doc, 2, "3.3 ChatDev（清华/OpenBMB）——轻量开源多智能体")
para(doc, "概况：开源（Apache 2.0），以 chat-chain 沿瀑布阶段（设计→编码→测试→文档）"
     "组织角色对话，提出「交流去幻觉」机制，学术对比中质量指标优于部分单智能体方案。"
     "局限与 MetaGPT 类似：偏研究验证，缺少成本治理、执行沙箱、产品化形态与"
     "商业支持——token-burner 在这些维度均为工程化补全。")

h(doc, 2, "3.4 AI 编码助手三强：Claude Code / Cursor / GitHub Copilot")
table(doc, [
    ("产品", "定位与价格", "关键数据", "与 token-burner 的关系"),
    ("Claude Code", "终端优先自主代理，$20-200/月", "SWE-bench Verified 80.8%"
     "（2026 年最高），1M 上下文", "单模型代理能力最强，但单厂商锁定、"
     "无跨模型互验、无任务级预算治理"),
    ("Cursor", "AI 原生 IDE（VS Code 分叉），$20/月（旗舰版 $200/月）",
     "SWE-bench 约 51.7-65%，补全接受率约 72%",
     "IDE 内编辑体验最佳，服务于「写代码的人」而非端到端项目交付"),
    ("GitHub Copilot", "多 IDE 嵌入式助手，$10/月起", "470 万付费订阅、约 2000 万"
     "用户、Fortune 100 覆盖约 90%", "普及度最高，但单点补全/对话为主，"
     "无多智能体编排与预算工程"),
], widths=[1.1, 1.9, 1.9, 2.1])
para(doc, "小结：三者验证了「AI 编码」是商业化最快的 AI 赛道（Cursor ARR 达 20 亿美元"
     "量级、Copilot 订阅 470 万），但其产品形态均以「辅助开发者」为中心；"
     "token-burner 面向的是另一个买方诉求——「给我一个可运行的项目，且成本可预算、"
     "过程可审计」。两类产品在企业内是互补而非替代关系。")

h(doc, 2, "3.5 其他值得关注的玩家")
bullet(doc, "Microsoft AutoGen / CrewAI：通用多智能体编排框架，灵活但需自建"
       "软件工程流程与护栏——token-burner 相当于把这套流程开箱即用。")
bullet(doc, "Windsurf（$15/月）：Cursor 的平价替代，后被 Cognition 收购整合为 "
       "Devin Desktop。")
bullet(doc, "Aider：开源 Git 原生 CLI（自带密钥），轻量代理但对「项目级交付、"
       "多模型互验、预算治理」无覆盖。")

h(doc, 1, "四、横向对比矩阵")
table(doc, [
    ("维度", "token-burner", "Devin", "MetaGPT/ChatDev", "Copilot/Cursor/CC"),
    ("产品形态", "Web/桌面 exe/CLI/VS Code 插件", "云端 SaaS + IDE",
     "开源框架/库", "IDE 插件 / 终端"),
    ("多智能体", "三模型强制互异（PM+Dev+Test）", "单一代理（模型路由）",
     "多角色（5 角色/双人对话）", "单模型代理"),
    ("成本治理", "六层确定性护栏 + 200k 预算总闸 + 成本看板",
     "ACU 后付费计量", "无", "订阅额度/信用点"),
    ("开工前成本可预估", "是（token 上界固定）", "否（按 ACU 事后计）", "否", "否"),
    ("验证闭环", "静态验证+AST 接口门禁+双模执行+反馈闭环",
     "沙箱内自测 + PR", "生成测试文件", "开发者自测"),
    ("执行沙箱", "Docker 配额沙箱（可降级进程模式）", "云端 VM", "无",
     "本地/受限"),
    ("模型中立（BYOK）", "是（litellm 任意组合）", "否（厂商路由）", "是", "否（绑定厂商）"),
    ("私有化部署", "是（本地运行，数据不出域）", "VPC（企业版）", "是（自建）",
     "否"),
    ("过程可审计", "逐调用 token 落盘 + 对话流 + 修复记录 + 中断快照",
     "平台内日志", "无系统化审计", "无任务级审计"),
    ("开源/开放", "内核本地可审计", "闭源", "开源（MIT/Apache）", "闭源"),
    ("价格", "自带模型密钥（按用量）", "$20-500/月 + ACU", "免费（自付 API）",
     "$10-200/月"),
], widths=[1.3, 1.6, 1.3, 1.3, 1.5], size=8.5)

h(doc, 1, "五、token-burner 的差异化定位")
bullet(doc, "同类产品中唯一把成本治理做成「确定性机制」的：Devin 的成本控制是计费"
       "提醒（ACU 事后账单），MetaGPT 类框架几乎没有成本治理；token-burner 的六层"
       "护栏 + 预算总闸让任务开工前即锁定 token 上界、超支即熔断落盘——这是预算管理"
       "岗位（采购/财务/项目经理）的采购语言。", "1. 成本治理的代际差：")
bullet(doc, "三模型强制互异 + 双评审 + 收敛裁决，用模型异构性对冲单一模型盲区；"
       "Claude Code/Copilot 类单模型产品在结构上无法复制。", "2. 异构互验架构：")
bullet(doc, "相对 MetaGPT/ChatDev 等研究框架，token-burner 补齐了产品化所需的全部"
       "工程件：中断恢复、任务取消、异步 API+SSE、并发隔离、限流缓存、四端入口、"
       "发布流水线与 833 项测试。", "3. 工程化完整度：")
bullet(doc, "本地运行 + 密钥自带 + 模型任意组合：需求文本、生成代码、讨论过程均"
       "不出域，直接命中数据合规与供应商中立诉求——这是对 Devin 类 SaaS 的结构性"
       "差异。", "4. 私有化与模型中立：")
bullet(doc, "逐调用 token 记录、阶段/模型/档位成本分解、对话流图、修复历史、"
       "中断快照全量落盘，合规与审计部门可直接取证。", "5. 全链路审计：")
bullet(doc, "「决策归 LLM、校验归程序」的边界划分 + SDD/TDD 方法论 + 规格锚点，"
       "使产品质量可机检、迭代有回归安全网。", "6. 可验证的质量方法论：")

h(doc, 1, "六、客观差距与风险（诚实面对）")
table(doc, [
    ("差距/风险", "客观描述", "应对方向"),
    ("品牌与生态位", "Devin 已进入 Fortune 50 与联邦市场；Copilot 470 万付费订阅；"
     "本项目尚无市场认知", "以「成本治理 + 私有化」切入点差异化突围，"
     "优先获取对预算与数据敏感的客群"),
    ("公开基准缺失", "竞品有 SWE-bench（Claude Code 80.8%）等公开成绩，"
     "token-burner 尚无第三方基准数据", "规划在标准基准（如 SWE-bench 子集）上"
     "跑分并公开方法论与成本对照"),
    ("复杂任务成功率天花板", "行业现实：Devin 复杂多文件任务成功率仅 14-15%；"
     "编排层无法完全弥补基座模型差距", "以三模型互验+接口门禁+修复上限收敛质量，"
     "并如实向用户呈现降级方案而非掩盖失败"),
    ("自动模式真实成功率待验证", "833 项测试为 stub 回归，真实 LLM 全链路成功率"
     "需大样本统计", "利用内置 A/B 框架与性能基线脚本持续积累真实数据"),
    ("分发摩擦", "未签名 exe 会触发部分安全软件误报（已架构性规避，正式发布需"
     "代码签名证书）", "发布流水线已就绪，商业化前采购签名证书"),
], widths=[1.4, 3.1, 2.5])

h(doc, 1, "七、场景适配建议")
table(doc, [
    ("如果你的场景是…", "更合适的选择", "原因"),
    ("企业级云端全自主开发，接受数据出域与 ACU 计费", "Devin",
     "企业治理成熟、生态集成（GitHub/Slack）完善"),
    ("研究人员复现/改进多智能体协作算法", "MetaGPT / ChatDev",
     "开源可改、社区与论文基线完善"),
    ("开发者日常编码提效（补全/重构/问答）", "Copilot / Cursor / Claude Code",
     "IDE 体验与模型能力最强"),
    ("要项目交付 + 成本必须可预算 + 数据不出域 + 模型自由组合",
     "token-burner", "唯一同时满足四项约束的产品"),
], widths=[2.6, 1.7, 2.7])

h(doc, 1, "八、结论")
para(doc, "2026 年的 AI 编码市场已充分验证「辅助开发者」赛道（Copilot/Cursor），"
     "并在「全自主工程师」赛道快速商业化（Devin）；但两类产品共同留出的空白是："
     "成本开工前不可预估、过程不可本地审计、模型不可自由组合。token-burner 以"
     "「多智能体编排 + 确定性成本护栏 + 私有化部署」精确落在这个空白上——"
     "它不与 Copilot 拼编辑体验、不与 Devin 拼企业生态，而是用「预算工程」这一"
     "未被占领的价值层建立自己的生态位。")
para(doc, "核心结论：token-burner 的差异化是结构性的（架构决定，竞品难以复制），"
     "主要短板是市场性的（品牌、基准数据、真实样本量）——后者可通过聚焦种子客群、"
     "公开基准评测与持续积累真实任务数据系统性弥补。", bold=True)

h(doc, 1, "附录：主要信息来源")
srcs = [
    "Devin 2.0/2.2 评测与定价（ACU 计量、78% 与 14-15% 成功率）：thebestaitools.co，2026-07",
    "Devin 官方能力与 FedRAMP/企业客户信息：cognition.com、AWS Marketplace、devin.ai，2026",
    "MetaGPT 框架与「Code = SOP(Team)」理念、69k star、MGX/Atoms 商业化：aiwiki.ai、"
    "arXiv:2308.00352（ICLR 2024 Oral），2026-07 更新",
    "ChatDev chat-chain 与交流去幻觉机制：arXiv:2307.07924",
    "Claude Code / Cursor / Copilot 2026 对比（定价、SWE-bench、市场份额）："
    "toolchase.com、aiunpacker.com、brightcoding.dev，2026 年 5-8 月",
    "token-burner 产品事实：本项目代码库 v0.5.0-beta、CHANGELOG.md、README.md，2026-09",
]
for s in srcs:
    bullet(doc, s)

out2 = ROOT / "token-burner_竞品对比分析报告.docx"
doc.save(out2)
print("saved:", out2)
